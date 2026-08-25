#!/usr/bin/env python3
"""
STREAMLIT APP - ARBEITSVERTRAG + PERSONALFRAGEBOGEN
Alle Fragen für die gelben Felder
Auf Russisch 🇷🇺
"""

import streamlit as st
from docx import Document
import io
import zipfile
from datetime import datetime
import os

st.set_page_config(page_title="Генератор документов", page_icon="📋", layout="wide")

st.markdown("""
    <style>
    .title { font-size: 2.5em; color: #1F4E78; text-align: center; margin-bottom: 10px; }
    .subtitle { font-size: 1.2em; color: #666; text-align: center; margin-bottom: 30px; }
    .section { font-size: 1.2em; color: #1F4E78; border-bottom: 2px solid #D9E1F2; padding-bottom: 10px; margin-top: 20px; }
    .error { background-color: #F8D7DA; padding: 15px; border-radius: 5px; color: #721C24; }
    .success { background-color: #D4EDDA; padding: 15px; border-radius: 5px; color: #155724; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="title">📋 Генератор документов</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Трудовой договор + Анкета сотрудника</div>', unsafe_allow_html=True)

# FORM
st.markdown('<div class="section">👤 ЛИЧНЫЕ ДАННЫЕ</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    vorname = st.text_input("Имя *", help="Например: Дмитро")
with col2:
    nachname = st.text_input("Фамилия *", help="Например: Цимбал")

col1, col2 = st.columns(2)
with col1:
    nachname_geburt = st.text_input("Девичья фамилия", help="Если есть")
with col2:
    geburtsdatum = st.text_input("Дата рождения (ДД.MM.ГГГГ) *", help="Например: 15.03.1990")

col1, col2 = st.columns(2)
with col1:
    strasse = st.text_input("Улица и номер дома", help="Например: Иоганнисштр. 32")
with col2:
    plz = st.text_input("Почтовый индекс", help="Например: 08393")

stadt = st.text_input("Город", help="Например: Мерана")
nationalitaet = st.text_input("Национальность", help="Например: Украинская")

col1, col2 = st.columns(2)
with col1:
    familienstand = st.selectbox("Семейное положение", ["холост/холоста", "женат/замужем", "разведен/а", "вдовец/вдова"])
with col2:
    steuer_id = st.text_input("Налоговый номер", help="11-значный номер")

st.markdown('<div class="section">💼 УСЛОВИЯ РАБОТЫ</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    arbeitsbeginn = st.text_input("Дата начала (ДД.MM.ГГГГ) *", help="Например: 10.08.2026")
with col2:
    position = st.text_input("Должность *", value="Складской рабочий")

col1, col2, col3 = st.columns(3)
with col1:
    stundensatz = st.text_input("Почасовая ставка (€)", value="15,69 €")
with col2:
    arbeitszeit = st.selectbox("Тип работы", ["Vollzeit", "Teilzeit", "Minijob"])
with col3:
    gehalt = st.text_input("Зарплата", value="см. трудовой договор")

st.markdown('<div class="section">🏦 БАНКОВСКИЕ ДАННЫЕ</div>', unsafe_allow_html=True)

bankname = st.text_input("Банк", value="Sparkasse")

col1, col2 = st.columns(2)
with col1:
    bankleitzahl_praefix = st.text_input("BLZ Префикс", value="BE")
with col2:
    bankleitzahl_suffix = st.text_input("BLZ Суффикс", value="GTZU")

iban = st.text_input("IBAN", help="Например: DE04 2365 0236 1712 83 00")

st.markdown('<div class="section">📋 ДОПОЛНИТЕЛЬНО</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    krankenkasse = st.text_input("Страховая касса", value="AOK")
with col2:
    religionsgemeinschaft = st.selectbox("Религиозная община", ["ev", "kath", "andere", "keine"])

col1, col2 = st.columns(2)
with col1:
    datum_ausgefuellt = st.text_input("Дата заполнения (ДД.MM.ГГГГ)", value=datetime.now().strftime("%d.%m.%Y"))
with col2:
    vertragsbeginn = st.text_input("Начало контракта (ДД.MM.ГГГГ)", value="01.01.2026")

col1, col2 = st.columns(2)
with col1:
    personalennummer = st.text_input("Личный номер", value="LO2365")
with col2:
    steuer_info = st.text_input("Информация НДФЛ", value="02 020126 L 23")

ort = st.text_input("Город подписания", value="Berlin")
datum = st.text_input("Дата (ДД.MM.ГГГГ)", value=datetime.now().strftime("%d.%m.%Y"))

# BUTTON
st.markdown("---")
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    if st.button("✅ СОЗДАТЬ ДОКУМЕНТЫ", use_container_width=True):
        
        # VALIDATION
        errors = []
        if not vorname:
            errors.append("❌ Имя обязательно")
        if not nachname:
            errors.append("❌ Фамилия обязательна")
        if not geburtsdatum:
            errors.append("❌ Дата рождения обязательна")
        if not arbeitsbeginn:
            errors.append("❌ Дата начала обязательна")
        if not position:
            errors.append("❌ Должность обязательна")
        
        if errors:
            for error in errors:
                st.markdown(f'<div class="error">{error}</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="success">⏳ Создаю документы...</div>', unsafe_allow_html=True)
            
            # PREPARE DATA
            data = {
                "VORNAME": vorname,
                "NACHNAME": nachname,
                "NACHNAME_GEBURT": nachname_geburt,
                "GEBURTSDATUM": geburtsdatum,
                "STRASSE": strasse,
                "PLZ": plz,
                "STADT": stadt,
                "NATIONALITAET": nationalitaet,
                "FAMILIENSTAND": familienstand,
                "STEUER_ID": steuer_id,
                "ARBEITSBEGINN_TAG": arbeitsbeginn.split('.')[0] if arbeitsbeginn else "10",
                "ARBEITSBEGINN_MONAT": arbeitsbeginn.split('.')[1] if len(arbeitsbeginn.split('.')) > 1 else "08",
                "ARBEITSBEGINN_JAHR": arbeitsbeginn.split('.')[2] if len(arbeitsbeginn.split('.')) > 2 else "2026",
                "POSITION": position,
                "STUNDENSATZ": stundensatz,
                "ARBEITSZEIT": arbeitszeit,
                "GEHALT": gehalt,
                "BANKNAME": bankname,
                "BANKLEITZAHL_PRAEFIX": bankleitzahl_praefix,
                "BANKLEITZAHL_SUFFIX": bankleitzahl_suffix,
                "IBAN": iban,
                "KRANKENKASSE": krankenkasse,
                "RELIGIONSGEMEINSCHAFT": religionsgemeinschaft,
                "DATUM_AUSGEFUELLT": datum_ausgefuellt,
                "VERTRAGSBEGINN": vertragsbeginn,
                "PERSONALENNUMMER": personalennummer,
                "STEUER_INFO": steuer_info,
                "ORT": ort,
                "DATUM": datum,
            }
            
            # LOAD & REPLACE
            try:
                templates = [
                    ("Arbeitsvertrag_Vorlage.docx", "01_Arbeitsvertrag"),
                    ("Personalfragebogen_Vorlage.docx", "02_Personalfragebogen"),
                ]
                
                # Create ZIP
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    
                    doc_count = 0
                    
                    for template_file, doc_name in templates:
                        if os.path.exists(template_file):
                            doc = Document(template_file)
                            
                            # Replace in paragraphs
                            for paragraph in doc.paragraphs:
                                for run in paragraph.runs:
                                    for placeholder, value in data.items():
                                        if f"{{{{{placeholder}}}}}" in run.text:
                                            run.text = run.text.replace(f"{{{{{placeholder}}}}}", str(value))
                            
                            # Replace in tables
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                for placeholder, value in data.items():
                                                    if f"{{{{{placeholder}}}}}" in run.text:
                                                        run.text = run.text.replace(f"{{{{{placeholder}}}}}", str(value))
                            
                            # Save to ZIP
                            doc_buffer = io.BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            zip_file.writestr(f"{doc_name}_{nachname}_{vorname}.docx", 
                                           doc_buffer.getvalue())
                            doc_count += 1
                    
                    # SUCCESS
                    st.markdown(f'<div class="success">✅ Успешно! Создано {doc_count} документов</div>', 
                              unsafe_allow_html=True)
                    
                    # DOWNLOAD
                    zip_buffer.seek(0)
                    st.download_button(
                        label=f"📥 Скачать документы ({doc_count} шт)",
                        data=zip_buffer.getvalue(),
                        file_name=f"Dokumenty_{nachname}_{vorname}_{datetime.now().strftime('%Y%m%d')}.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
                    
            except Exception as e:
                st.markdown(f'<div class="error">❌ Ошибка: {str(e)}</div>', unsafe_allow_html=True)

# FOOTER
st.markdown("---")
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown("**EKO-ARK GmbH** © 2026")
with col2:
    st.markdown("📧 info@ekoark.de")
with col3:
    st.markdown("📱 +49 30 67384131")
